#!/usr/bin/env python3
"""
Module to generate lesson notes template for Morning Star School
"""
import os
import re
import logging
import tempfile
from typing import Dict, Any

import matplotlib
import matplotlib.pyplot as plt
import warnings
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from markdown2 import markdown
from bs4 import BeautifulSoup

# Suppress deprecated style_id warning
warnings.filterwarnings(
    "ignore", category=UserWarning, module="docx.styles.styles"
)

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# Directory to serve files from
UPLOAD_FOLDER = 'generated_files'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)


def set_cell_text(cell, text, bold=False, font_size=12, align='center'):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = getattr(WD_ALIGN_PARAGRAPH, align.upper())
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def render_latex_to_image(latex_text: str, output_path: str) -> bool:
    try:
        matplotlib.use('Agg')
        plt.figure(figsize=(8, 2), dpi=200)
        plt.rc('text', usetex=True)
        plt.rc('font', family='serif', size=14)
        plt.rc('text.latex', preamble=r'\usepackage{amsmath}\usepackage{amssymb}')
        plt.text(0.5, 0.5, f"${latex_text}$", ha='center', va='center')
        plt.axis('off')
        plt.savefig(output_path, bbox_inches='tight', pad_inches=0.05, transparent=True)
        plt.close()
        logger.debug(f"LaTeX rendered: {latex_text}")
        return True
    except Exception as e:
        logger.warning(f"LaTeX failed for '{latex_text}': {e}, falling back to mathtext.")
        try:
            plt.rc('text', usetex=False)
            plt.figure(figsize=(8, 2), dpi=200)
            plt.rc('font', family='serif', size=14)
            plt.text(0.5, 0.5, f"${latex_text}$", ha='center', va='center')
            plt.axis('off')
            plt.savefig(output_path, bbox_inches='tight', pad_inches=0.05, transparent=True)
            plt.close()
            logger.debug(f"Mathtext rendered: {latex_text}")
            return True
        except Exception as e2:
            logger.error(f"Mathtext failed for '{latex_text}': {e2}")
            return False


def detect_and_process_latex(cell, text: str):
    latex_pattern = r"\$(.*?)\$|\\\((.*?)\\\)|\\\[([\s\S]*?)\\\]"
    parts = []
    last_end = 0
    for match in re.finditer(latex_pattern, text, re.DOTALL):
        start, end = match.span()
        if last_end < start:
            parts.append(('text', text[last_end:start]))
        content = next(g for g in match.groups() if g)
        content = re.sub(r'\\{2,}', r'\\', content).strip()
        parts.append(('latex', content))
        last_end = end
    if last_end < len(text):
        parts.append(('text', text[last_end:]))

    for typ, cont in parts:
        if typ == 'text':
            # Split double newlines into paragraphs
            paras = re.split(r'\n{2,}', cont)
            for para in paras:
                if not para.strip():
                    continue
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lines = para.split('\n')
                for i, line in enumerate(lines):
                    add_markdown_to_paragraph(cell, line, paragraph=p)
                    if i < len(lines) - 1:
                        p.add_run().add_break()
        else:
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                if render_latex_to_image(cont, tmp.name):
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(tmp.name, width=Inches(3.5))
                    os.unlink(tmp.name)
                else:
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run(cont)


def apply_html_styles(cell, html_content, paragraph=None):
    soup = BeautifulSoup(html_content, 'html.parser')
    def recurse(node, p, bold, ital, under):
        if node.name == 'p':
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for c in node.children:
                recurse(c, p, bold, ital, under)
        elif node.name:
            nb = bold or node.name in ('b','strong')
            ni = ital or node.name in ('i','em')
            nu = under or node.name=='u'
            for c in node.children:
                recurse(c, p, nb, ni, nu)
        elif node.string and node.string.strip():
            run = p.add_run(node.string)
            run.bold, run.italic, run.underline = bold, ital, under
            run.font.name='Times New Roman'; run.font.size=Pt(14)
    base_para = paragraph or cell.paragraphs[0]
    for child in soup.contents:
        recurse(child, base_para, False, False, False)


def add_markdown_to_paragraph(cell, text, paragraph=None):
    if re.search(r"\$(.*?)\$", text):
        detect_and_process_latex(cell, text)
        return
    html = markdown(text) if '<' not in text else text
    apply_html_styles(cell, html, paragraph)


def add_bulleted_list(cell, items):
    """
    Add a bulleted list to a cell, with each item supporting HTML/Markdown styling
    """
    for item in items:
        p = cell.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        temp_doc = Document()
        temp_cell = temp_doc.add_table(1, 1).rows[0].cells[0]
        add_markdown_to_paragraph(temp_cell, item.strip())
        for temp_para in temp_cell.paragraphs:
            for run in temp_para.runs:
                new_run = p.add_run(run.text)
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline

def add_paragraphs_to_cell(cell, text):
    # Use unified latex+newline handler
    detect_and_process_latex(cell, text)

    
def sanitize_filename(filename):
    """
    Sanitize filename to remove invalid characters
    """
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

def create_lesson_notes_template(data=None, logo_path='./assets/images/MostarLogo.png'):
    """
    Creates a lesson notes template for Morning Star School and returns the file path
    """
    if data is None:
        data = {}
    
    doc = Document()

    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    PAGE_WIDTH = section.page_width - Inches(2)

    try:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("THE MORNING STAR SCHOOL LTD.\n")
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = 'Times New Roman'

        run.add_break()
        if os.path.exists(logo_path):
            run.add_picture(logo_path, width=Inches(1.5))
        else:
            logger.warning(f"Logo file not found at {logo_path}")
        run.add_break()
        run.add_text("WEEKLY LESSON PLAN")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        logger.error(f"Error adding header: {e}")
        raise

    doc.add_paragraph()

    rows_data = [
        ("WEEK ENDING", data.get("WEEK_ENDING", "")),
        ("DAYS", " ".join(data.get("DAYS", [])) if isinstance(data.get("DAYS"), list) else data.get("DAYS", "")),
        ("DURATION", data.get("DURATION", "")),
        ("SUBJECT", data.get("SUBJECT", "")),
        ("STRAND", data.get("STRAND", "")),
        ("SUBSTRAND", data.get("SUBSTRAND", "")),
        ("CLASS", data.get("CLASS", "")),
        ("CLASS SIZE", " ".join(f"{cls}({size})" for cls, size in data.get("CLASS_SIZE", {}).items()) if isinstance(data.get("CLASS_SIZE"), dict) else data.get("CLASS_SIZE", "")),
        ("CONTENT STANDARD (ANNOTATION)", data.get("CONTENT_STANDARD", [])),
        ("LEARNING INDICATOR(S)", data.get("LEARNING_INDICATORS", [])),
        ("PERFORMANCE INDICATOR(S)", data.get("PERFORMANCE_INDICATORS", [])),
        ("TEACHING/LEARNING RESOURCES (TLMS)", data.get("TEACHING_LEARNING_RESOURCES", [])),
        ("CORE COMPETENCIES", data.get("CORE_COMPETENCIES", [])),
        ("KEY WORDS", data.get("KEY_WORDS", [])),
        ("R.P.K", data.get("R.P.K", "")),
    ]

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    table.autofit = False

    col1_width = PAGE_WIDTH * 0.3
    col2_width = PAGE_WIDTH * 0.7

    for i, (label, value) in enumerate(rows_data):
        cell1, cell2 = table.rows[i].cells
        cell1.width = col1_width
        cell2.width = col2_width
        set_cell_text(cell1, label, bold=True, font_size=16)
        if isinstance(value, list):
            add_bulleted_list(cell2, value)
        else:
            add_paragraphs_to_cell(cell2, value)

    doc.add_paragraph()

    phase_headers = ["PHASE 1: STARTER", "PHASE 2: MAIN", "PHASE 3: REFLECTION"]
    phase_data = (
        data.get("PHASE_1", {}).get("STARTER", "") if isinstance(data.get("PHASE_1"), dict) else data.get("PHASE_1", ""),
        data.get("PHASE_2", {}).get("MAIN", "") if isinstance(data.get("PHASE_2"), dict) else data.get("PHASE_2", ""),
        data.get("PHASE_3", {}).get("REFLECTION", "") if isinstance(data.get("PHASE_3"), dict) else data.get("PHASE_3", "")
    )

    table2 = doc.add_table(rows=2, cols=3)
    table2.style = 'Table Grid'
    table2.autofit = False

    col1_width = PAGE_WIDTH * 0.2
    col2_width = PAGE_WIDTH * 0.6
    col3_width = PAGE_WIDTH * 0.2

    for i, (row_data) in enumerate([phase_headers, phase_data]):
        cell1, cell2, cell3 = table2.rows[i].cells
        cell1.width = col1_width
        cell2.width = col2_width
        cell3.width = col3_width
        if i == 0:
            set_cell_text(cell1, row_data[0], bold=True, font_size=14, align="justify")
            set_cell_text(cell2, row_data[1], bold=True, font_size=14, align="center")
            set_cell_text(cell3, row_data[2], bold=True, font_size=14, align="center")
        else:
            add_paragraphs_to_cell(cell1, row_data[0])
            add_paragraphs_to_cell(cell2, row_data[1])
            add_paragraphs_to_cell(cell3, row_data[2])

    doc.add_paragraph()

    table3 = doc.add_table(rows=2, cols=2)
    table3.style = 'Table Grid'
    table3.autofit = False
    col1_width = PAGE_WIDTH * 0.3
    col2_width = PAGE_WIDTH * 0.7

    assessment_data = [
        ("ASSESSMENTS", ""),
        ("", f"{data.get('ASSESSMENTS', '')}\n\n\n{data.get('HOMEWORK', '')}"),
    ]

    for i, (label, content) in enumerate(assessment_data):
        cell1, cell2 = table3.rows[i].cells
        cell1.width = col1_width
        cell2.width = col2_width
        set_cell_text(cell1, label, bold=True, font_size=16)
        add_paragraphs_to_cell(cell2, content)

    subject = data.get("SUBJECT", "Unknown")
    week = data.get("WEEK", "Unknown")
    cls = data.get("CLASS", "Unknown")
    filename = sanitize_filename(f"{cls} Lesson Notes {subject} WEEK {week}.docx")
    try:
        doc.save(filename)
        file_path = os.path.abspath(filename)
        logger.info(f"Lesson notes template created successfully: {file_path}")
        return file_path
    except Exception as e:
        logger.error(f"Error saving document: {e}")
        raise


if __name__ == "__main__":
    example_data = {
    "WEEK_ENDING": "16th May, 2025",
    "DAYS": "Monday - Friday",
    "WEEK": "3",
    "DURATION": "4 periods per class",
    "SUBJECT": "Mathematics",
    "STRAND": "Strand 3: Geometry and Measurement",
    "SUBSTRAND": "Substrand 2: Angles and Polygons",
    "CLASS": "Basic Eight",
    "CLASS_SIZE": {"A": 28, "B": 28, "C": 28},
    "CONTENT_STANDARD": ["B8.3.2.1: Demonstrate understanding of properties of polygons and solve related problems"],
    "LEARNING_INDICATORS": ["B8.3.2.1.1: Identify and calculate interior and exterior angles of polygons"],
    "PERFORMANCE_INDICATORS": [
        "Calculate the sum of interior angles of various polygons.",
        "Determine the measure of an exterior angle of a regular polygon.",
        "Solve real-life problems involving angles in polygons."
    ],
    "TEACHING_LEARNING_RESOURCES": ["Charts showing different polygons", "Markers", "Whiteboard", "Protractor"],
    "CORE_COMPETENCIES": ["Creativity", "Critical Thinking", "Collaboration"],
    "KEY_WORDS": ["Polygon", "Interior Angle", "Exterior Angle", "Regular Polygon", "Irregular Polygon", "Sum of Angles", "Vertex"],
    "R.P.K": "Learners have basic knowledge of triangles and quadrilaterals and can identify polygons with up to six sides.",
    "PHASE_1": {
        "STARTER": "Begin the lesson by asking students to name different shapes they see around them and classify them as polygons or non-polygons. Discuss what makes a shape a polygon and introduce the concept of angles in these shapes."
    },
    "PHASE_2": {
        "MAIN": "The objective of this lesson is to understand and calculate the angles in various polygons, which are essential components in both mathematics and everyday life. We will explore both interior and exterior angles and apply these concepts to solve problems. \n\n1. **Lesson Objective:** By the end of the lesson, learners should be able to calculate the sum of interior angles and the measure of exterior angles in polygons. \n\n2. **Introduction:** Consider the Ghanaian Kente cloth, which often features geometric patterns. These patterns include various polygons, such as triangles, squares, and hexagons. Understanding the properties of these shapes helps in creating precise and beautiful designs. Similarly, the architecture of traditional Ghanaian buildings often incorporates polygonal shapes for aesthetic and structural purposes. \n\n3. **Step-by-Step Explanation:** \n   - **Interior Angles:** The sum of the interior angles of a polygon with \( n \) sides can be calculated using the formula: \n     \[ (n - 2) \times 180^\circ \] \n     For example, a pentagon (5 sides) has an interior angle sum of: \n     \[ (5 - 2) \times 180^\circ = 540^\circ \] \n   - **Exterior Angles:** The sum of the exterior angles of any polygon is always \( 360^\circ \). For a regular polygon, each exterior angle can be calculated by dividing \( 360^\circ \) by the number of sides \( n \): \n     \[ \text{Exterior Angle} = \frac{360^\circ}{n} \] \n     For a hexagon, each exterior angle is: \n     \[ \frac{360^\circ}{6} = 60^\circ \] \n\n4. **Guided Practice:** \n   - **Activity 1:** Use the protractor to measure angles of different polygons drawn on the whiteboard. Collaboratively calculate the interior angle sum for each shape. \n   - **Activity 2:** In groups, create and decorate a polygon cut-out (triangle, square, pentagon) and label each angle. Calculate both interior and exterior angles. \n\n5. **Independent Practice:** \n   - Problem 1: Calculate the sum of interior angles of a nonagon (9 sides). \n   - Problem 2: Determine the measure of one interior angle of a regular octagon. \n   - Problem 3: A polygon has an exterior angle of \( 45^\circ \). How many sides does this polygon have?"
    },
    "PHASE_3": {
        "REFLECTION": "Review the key concepts of interior and exterior angles in polygons. Ask students to share how they might use these calculations in real-life scenarios, such as designing patterns or constructing objects. Clarify any mistakes and emphasize the importance of accuracy in calculations. Discuss how understanding these concepts can aid in solving more complex geometrical problems."
    },
    "ASSESSMENTS": "Observe learners as they engage in activities, ensuring they collaborate effectively and understand the concepts. Provide immediate feedback and address any misconceptions during the lesson. Use a short quiz at the end to assess their understanding of interior and exterior angles in polygons.",
    "HOMEWORK": "Complete the following problems: 1) Calculate the sum of the interior angles of a decagon (10 sides). 2) If each exterior angle of a regular polygon is \( 30^\circ \), how many sides does the polygon have? Ensure to show all your workings."
}
    create_lesson_notes_template(example_data)
